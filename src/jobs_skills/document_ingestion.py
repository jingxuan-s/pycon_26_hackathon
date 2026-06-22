"""Runtime document ingestion for local resume/JD workflows.

The ingestion layer reads source documents into memory only. Persisted artifacts must
store extracted skill evidence and reports, not raw resume or JD text.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass(frozen=True)
class ExtractedDocument:
    source_path: Path
    source_type: str
    text: str
    extraction_notes: tuple[str, ...]


def extract_text_from_file(path: str | Path) -> ExtractedDocument:
    """Extract text from a PDF, DOCX, TXT, or MD file without persisting it."""
    source_path = Path(path).expanduser().resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"Document not found: {source_path}")
    if not source_path.is_file():
        raise ValueError(f"Document path is not a file: {source_path}")

    extension = source_path.suffix.casefold()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
        raise ValueError(f"Unsupported document type {extension!r}. Supported: {supported}")

    if extension == ".pdf":
        text, notes = _extract_pdf(source_path)
    elif extension == ".docx":
        text, notes = _extract_docx(source_path)
    else:
        text = source_path.read_text(encoding="utf-8", errors="replace")
        notes = ("Plain text document read with UTF-8 replacement for invalid bytes.",)

    cleaned = _clean_text(text)
    if not cleaned:
        notes = tuple(notes) + ("No extractable text was found.",)
    return ExtractedDocument(
        source_path=source_path,
        source_type=extension.lstrip("."),
        text=cleaned,
        extraction_notes=tuple(notes),
    )


def document_text_from_pasted_input(text: str, source_type: str = "pasted_text") -> ExtractedDocument:
    """Wrap pasted JD/resume text in the same runtime-only document contract."""
    cleaned = _clean_text(text)
    notes = ("Pasted text was used at runtime and was not written as a raw source document.",)
    return ExtractedDocument(
        source_path=Path("<pasted>"),
        source_type=source_type,
        text=cleaned,
        extraction_notes=notes,
    )


def _extract_pdf(path: Path) -> tuple[str, tuple[str, ...]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)
        else:
            pages.append(f"\n[Page {index}: no extractable text]\n")
    return "\n".join(pages), (f"Extracted text from {len(reader.pages)} PDF page(s).",)


def _extract_docx(path: Path) -> tuple[str, tuple[str, ...]]:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts), ("Extracted text from DOCX paragraphs and tables.",)


def _clean_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()
