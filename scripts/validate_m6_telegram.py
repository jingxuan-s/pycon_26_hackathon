"""Validate Telegram v2 resume/role-first integration acceptance checks."""

from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC = PROJECT_ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from jobs_skills.scoring import select_role_id
from jobs_skills.telegram_bot import TelegramCareerBotService


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def button_actions(response) -> set[str]:
    return {button.action for button in response.buttons}


def button_labels(response) -> list[str]:
    return [button.label for button in response.buttons]


def first_button_value(response, action: str) -> str:
    for button in response.buttons:
        if button.action == action:
            return button.value
    raise AssertionError(f"Missing button action: {action}")


def attachment_text(response) -> str:
    if response.attachment_content is not None:
        content = response.attachment_content
        if isinstance(content, bytes):
            return content.decode("utf-8")
        return content
    if response.attachment_path:
        return Path(response.attachment_path).read_text(encoding="utf-8")
    raise AssertionError("Expected an attachment")


def main() -> int:
    os.environ["EXPLAINER_AGENT_DISABLED"] = "1"
    os.environ["PARSER_AGENT_DISABLED"] = "1"
    service = TelegramCareerBotService(PROJECT_ROOT)

    start = service.start("v2-start")
    require("Casey the Career Auntie" in start.text, "/start must show Casey v2 entry menu")
    require("Upload resume" in start.text and "Search role" in start.text, "/start must mention resume and role search")
    require(button_actions(start) == {"first_time_user", "start_resume", "start_role_search"}, "/start must expose v2 buttons only")
    require("Question 1/10" not in start.text, "/start must not launch the old questionnaire")

    guide = service.first_time_user("v2-start")
    require("First-time guide" in guide.text and "Privacy boundary" in guide.text, "First-time guide must explain workflow and privacy boundary")
    require({"start_resume", "start_role_search", "start"}.issubset(button_actions(guide)), "First-time guide must provide entry and back buttons")

    role_search = service.search_roles_entry("role-start", "data analyst")
    require("Role search" in role_search.text and "data analyst" in role_search.text, "Role search must return v2 role-start results")
    require("resume_start_from_role" in button_actions(role_search), "Role-start search must seed a skill profile")
    role_profile = service.start_profile_from_role("role-start", first_button_value(role_search, "resume_start_from_role"))
    require("draft skill profile" in role_profile.text, "Role selection must create a draft skill profile")
    require("\nReview all skills" in role_profile.text, "Remaining-skills prompt must be separated from top matches")
    require(button_actions(role_profile) == {"resume_review_page", "resume_explain_skills"}, "Draft profile must require review before scoring")

    review = service.review_resume_skills("role-start", 0)
    require("Review skills" in review.text and "1-5" in review.text, "Skill review must page through all skills")
    require("<b>Edit</b>" in review.text and "<b>Remove</b>" in review.text, "Review instructions must highlight edit/remove actions")
    require("Example: <b>Edit 1</b>" in review.text, "Review instructions must explain numbered edit buttons")
    require("confidence" not in review.text.casefold(), "Normal review must hide confidence")
    require("mapping:" not in review.text.casefold(), "Normal review must hide mapping internals")
    require({"resume_edit_skill", "resume_remove_skill", "resume_add_skill", "resume_confirm_skills", "resume_confirm_back_start"}.issubset(button_actions(review)), "Review page must expose edit/remove/add/continue/guarded-back")

    edit = service.edit_resume_skill("role-start", "0")
    require("Set level" in edit.text and "Level guide" in edit.text and "Dataset guidance" in edit.text and {"resume_set_level", "resume_back_review"}.issubset(button_actions(edit)), "Edit flow must ask for a level, show dataset guidance, and provide back navigation")
    require("Dataset rows available" in edit.text, "Edit flow must explain when only some dataset levels exist")
    require("Level 6" not in button_labels(edit), "Edit flow must hide unavailable K&A levels")
    updated = service.set_resume_skill_level("role-start", "0:2")
    require("Skill level updated" in updated.text, "Level update must return to review")
    require("(edited)" in updated.text, "Edited skills must be visibly labelled in review")
    before_count = len(service.get_session("role-start").resume_profile.items)
    removed = service.remove_resume_skill("role-start", "1")
    after_count = len(service.get_session("role-start").resume_profile.items)
    require(after_count == before_count - 1, "Remove skill must change the profile")
    require("Skill removed" in removed.text, "Remove skill must return to review")

    back_confirm = service.confirm_back_to_start("role-start")
    require("Back to start?" in back_confirm.text and {"start", "resume_back_review"}.issubset(button_actions(back_confirm)), "Back to start must require explicit confirmation")
    require("restart the workflow" in back_confirm.text, "Back to start confirmation must explain workflow restart")

    add_prompt = service.prompt_add_resume_skill("role-start")
    require("Type a skill name" in add_prompt.text and "resume_back_review" in button_actions(add_prompt), "Add skill must prompt for search text and provide back navigation")
    modelling_search = service.handle_text_message("role-start", "computational modeling")
    require("Computational Modelling" in modelling_search.text, "Skill search must tolerate modeling/modelling spelling variants")
    modelling_select = first_button_value(modelling_search, "resume_add_select")
    modelling_level = service.select_resume_add_skill("role-start", modelling_select)
    require(
        {"Level 0", "Level 3", "Level 4", "Level 5", "Back to review"} == set(button_labels(modelling_level)),
        "Add-skill level buttons must use K&A-backed levels for Computational Modelling",
    )
    service.prompt_add_resume_skill("role-start")
    add_search = service.handle_text_message("role-start", "programming")
    require("Skill search: programming" in add_search.text, "Add skill search must search SkillsFuture skills")
    add_select_value = first_button_value(add_search, "resume_add_select")
    add_level = service.select_resume_add_skill("role-start", add_select_value)
    require("Choose your current level" in add_level.text, "Add skill must ask for proficiency level")
    added = service.add_resume_skill_level("role-start", f"{add_select_value}:3")
    require("Skill added" in added.text, "Add skill must return to review")
    require("(added)" in added.text, "Added skills must be visibly labelled in review")

    confirmed = service.confirm_resume_skills("role-start")
    require("Casey can compare your profile" in confirmed.text and "higher-level roles" in confirmed.text, "Confirmed profile must describe target buttons")
    require("Already have a role in mind" in confirmed.text and "Applying for a job" in confirmed.text, "Confirmed profile must describe role search and JD buttons")
    for action in {"resume_explore", "resume_advance", "start_role_search", "resume_jd_text", "resume_back_review"}:
        require(action in button_actions(confirmed), f"Confirmed profile missing target action: {action}")

    explore = service.start_resume_explore("role-start")
    require("gaps" not in explore.text.casefold(), "Explore recommendations must hide gap counts until a target is selected")
    require("resume_back_targets" in button_actions(explore), "Explore recommendations must offer back navigation")
    require(len([button for button in explore.buttons if button.action == "resume_choose_role"]) >= 4, "Explore recommendations should show more than three target options when available")

    target_prompt = service.search_roles_entry("role-start", "")
    require("Type the role or domain in your next message" in target_prompt.text, "Target role search must accept plain next-message input")
    specific_search = service.handle_text_message("role-start", "data scientist")
    require("Target role search" in specific_search.text, "Role search after profile confirmation must search target roles")
    target_result = service.choose_resume_specific_role("role-start", first_button_value(specific_search, "resume_specific_role"))
    require("Suitability:" in target_result.text and "Top gaps" in target_result.text, "Specific target result must show compact suitability")
    require("parser source" not in target_result.text.casefold(), "Compact result must hide parser internals")
    require({"explain_score", "show_gaps", "resume_report", "generate_action_plan", "resume_back_targets"}.issubset(button_actions(target_result)), "Result must expose explain/gaps/report/action/back buttons")
    result_labels = button_labels(target_result)
    require(result_labels.index("Generate action plan") < result_labels.index("Generate report"), "Action plan button must appear before full report")
    require("<b>Why this score?</b>" in target_result.text, "Result text must highlight the Why this score button")
    require("resume_debug_report" not in button_actions(target_result), "Debug report button must be hidden by default")
    explain = service.explain_score("role-start")
    gaps = service.show_gaps("role-start")
    action_plan = service.generate_action_plan("role-start")
    require("Why this score?" in explain.text and "You match" in explain.text, "Why this score must explain active resume result")
    require("How it is calculated" not in explain.text, "Normal explanation must hide formula details")
    require("Top gaps" in gaps.text and "Source row" not in gaps.text, "Gaps must be compact and hide source rows")
    require(action_plan.attachment_content is not None and action_plan.attachment_name == "career_action_plan.md", "Normal action plan must generate an in-memory attachment")
    action_text = attachment_text(action_plan)
    require("### 1." in action_text, "Action plan must use separated markdown sections")
    require("K&A focus" in action_text or "Target proficiency" in action_text, "Action plan must include K&A/proficiency detail")

    related_service = TelegramCareerBotService(PROJECT_ROOT)
    related_source_role = select_role_id(
        related_service.requirements,
        "Data Scientist",
        sector="Financial Services",
        track="Digital and Data Analytics",
    )
    related_target_role = select_role_id(
        related_service.requirements,
        "Data Scientist / Artificial Intelligence Scientist",
        sector="Infocomm Technology",
        track="Data and Artificial Intelligence",
    )
    related_service.start_profile_from_role("related-skills", related_source_role)
    related_service.confirm_resume_skills("related-skills")
    related_result = related_service.choose_resume_specific_role("related-skills", related_target_role)
    related_score = related_service.get_session("related-skills").resume_result.selected_summary.suitability_percentage
    related_gaps = related_service.show_gaps("related-skills")
    related_explain = related_service.explain_score("related-skills")
    related_plan = related_service.generate_action_plan("related-skills")
    related_plan_text = attachment_text(related_plan)
    related_normal = related_service.generate_resume_report("related-skills", debug=False)
    related_debug = related_service.generate_resume_report("related-skills", debug=True)
    related_normal_text = attachment_text(related_normal)
    related_debug_text = attachment_text(related_debug)
    require("Computational Modelling" in related_gaps.text, "Detailed gaps should keep exact Computational Modelling gap visible")
    require("Related evidence" in related_gaps.text and "Data Analytics and Computational Modelling" in related_gaps.text, "Gaps must explain related but non-scored skills")
    require("Related evidence" in related_explain.text, "Why this score must include related-skill notes")
    require("Related evidence" in related_plan_text, "Action plan must include related-skill notes")
    require("Related Skill Notes" in related_normal_text, "Normal report must include related-skill explanation notes")
    require("Related-skill rule" in related_debug_text, "Debug report must include related-skill matching rule")
    require(related_service.get_session("related-skills").resume_result.selected_summary.suitability_percentage == related_score, "Related-skill notes must not change the suitability score")

    resume_doc = PROJECT_ROOT / ".runtime" / "telegram_validation_resume.docx"
    resume_doc.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("Sample Candidate Singaporean +65 8000 0000 LinkedIn candidate@example.com")
    doc.add_paragraph("Built Python scripts, SQL dashboards, and data pipelines for analytics projects.")
    doc.add_paragraph("Led process improvement work with business users and communicated findings.")
    doc.save(resume_doc)

    resume_start = service.start_resume_flow("resume-v2")
    require("Upload your resume" in resume_start.text, "Resume shortcut must ask for upload")
    resume_response = service.parse_resume_upload("resume-v2", resume_doc)
    require("Casey found" in resume_response.text, "Resume upload must return compact skill count")
    require("Review skills" in resume_response.text and "Review and confirm" in resume_response.text, "Resume upload must explicitly require the Review skills button")
    require("confidence" not in resume_response.text.casefold(), "Compact resume summary must hide confidence")
    require("mapping:" not in resume_response.text.casefold(), "Compact resume summary must hide mapping internals")
    require("8000 0000" not in resume_response.text and "candidate@example.com" not in resume_response.text.casefold(), "Compact resume summary must not expose contact details")
    require(button_actions(resume_response) == {"resume_review_page", "resume_explain_skills"}, "Resume summary must expose review buttons")
    why_skills = service.explain_resume_skills("resume-v2")
    require("<b>Why these skills?</b>" in why_skills.text, "Skill explainer header must be bolded")
    require("Evidence:" not in why_skills.text and "confidence" not in why_skills.text.casefold(), "Skill explainer should summarize without raw evidence labels")

    service.confirm_resume_skills("resume-v2")
    jd_prompt = service.start_resume_jd_text("resume-v2")
    require("Paste the job description" in jd_prompt.text, "Resume JD mode must ask for pasted JD text")
    jd_result = service.handle_text_message(
        "resume-v2",
        "We need a data scientist who can build Python models, SQL pipelines, dashboards, and communicate with business users.",
    )
    require("Suitability:" in jd_result.text, "Resume JD result must show suitability")
    require("Top gaps" in jd_result.text, "Resume JD result must show compact gaps")
    require("parser source" not in jd_result.text.casefold(), "Compact result must hide parser internals")
    jd_explain = service.explain_score("resume-v2")
    jd_gaps = service.show_gaps("resume-v2")
    require("No pathway selected" not in jd_explain.text + jd_gaps.text, "Resume/JD explain and gaps must not require old pathway selection")
    normal_report = service.generate_resume_report("resume-v2", debug=False)
    debug_report = service.generate_resume_report("resume-v2", debug=True)
    debug_service = TelegramCareerBotService(PROJECT_ROOT, debug_mode=True)
    debug_service.start_profile_from_role("debug-buttons", first_button_value(debug_service.search_roles_entry("debug-buttons", "data analyst"), "resume_start_from_role"))
    debug_service.confirm_resume_skills("debug-buttons")
    debug_target = debug_service.choose_resume_specific_role("debug-buttons", first_button_value(debug_service.search_roles_entry("debug-buttons", "data scientist"), "resume_specific_role"))
    require("resume_debug_report" in button_actions(debug_target), "Debug report button must appear when debug_mode is enabled")
    require(normal_report.attachment_content is not None and normal_report.attachment_path is None, "Normal report must be generated in memory without a local file")
    require(debug_report.attachment_path and Path(debug_report.attachment_path).exists(), "Debug report must be generated")
    normal_text = attachment_text(normal_report)
    debug_text = attachment_text(debug_report)
    require("Parser Diagnostics" not in normal_text and "Evidence snippet" not in normal_text, "Normal report must hide debug evidence")
    require("All MVP skill weights are 1.0" not in normal_text, "Normal report must hide MVP weight assumption")
    require("Parser Diagnostics" in debug_text and "Evidence snippet" in debug_text, "Debug report must include parser diagnostics")
    require("All MVP skill weights are 1.0" in debug_text, "Debug report must keep MVP weight assumption")

    print("M6 Telegram v2 validation passed")
    print(f"role_profile_skills={len(service.get_session('role-start').resume_profile.items)}")
    print(f"resume_profile_skills={len(service.get_session('resume-v2').resume_profile.items)}")
    print(f"action_plan={action_plan.attachment_name or action_plan.attachment_path}")
    print(f"normal_report={normal_report.attachment_name or normal_report.attachment_path}")
    print(f"debug_report={debug_report.attachment_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
